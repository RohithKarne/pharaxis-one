from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x52, 0x76)   # primary blue
TEAL   = RGBColor(0x17, 0xa5, 0x89)   # accent teal
DARK   = RGBColor(0x1c, 0x1c, 0x1c)   # body text
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)   # row shading

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ──────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'BBBBBB'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: paragraph style ───────────────────────────────────────────────────
def para(text, bold=False, size=11, colour=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(size)
    run.font.color.rgb = colour
    return p

# ── Helper: section heading ───────────────────────────────────────────────────
def section_heading(number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{number}. {title}")
    run.bold           = True
    run.font.size      = Pt(13)
    run.font.color.rgb = NAVY

# ── Helper: sub heading ───────────────────────────────────────────────────────
def sub_heading(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(title)
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = TEAL

# ── Helper: bullet ────────────────────────────────────────────────────────────
def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + level*0.2)
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = DARK

# ── Helper: styled table ──────────────────────────────────────────────────────
def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1A5276')
        set_cell_border(cell, color='1A5276')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg  = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, color='CCCCCC')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(val))
            run.font.size      = Pt(9)
            run.font.color.rgb = DARK

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacer


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("MIMS-CP PORTAL")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

para("Medical Information Management System + Customer Portal",
     bold=False, size=14, colour=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)

para("PMO HANDOVER DOCUMENT",
     bold=True, size=16, colour=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)

para("Sprint 0 → Sprint 3 | Project Status as of 11 March 2026",
     size=11, colour=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=60)

# Cover table
cover_data = [
    ("Document Type",    "PMO Handover Summary"),
    ("Version",          "Sprint 3 Close"),
    ("Prepared By",      "Claude (AI Testing Lead)"),
    ("Reviewed By",      "Rohith (CPO)"),
    ("Date",             "11 March 2026"),
    ("Classification",   "Internal — Confidential"),
]
ct = doc.add_table(rows=len(cover_data), cols=2)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
ct.style = 'Table Grid'
for i, (k, v) in enumerate(cover_data):
    bg = 'EAF2F8' if i % 2 == 0 else 'FFFFFF'
    lc = ct.rows[i].cells[0]; rc = ct.rows[i].cells[1]
    set_cell_bg(lc, '1A5276'); set_cell_bg(rc, bg)
    set_cell_border(lc, color='1A5276'); set_cell_border(rc, color='CCCCCC')
    lc.width = Inches(2.2); rc.width = Inches(3.8)
    p1 = lc.paragraphs[0]; p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(k); r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = WHITE
    p2 = rc.paragraphs[0]; p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(v); r2.font.size = Pt(10); r2.font.color.rgb = DARK

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
section_heading("1", "Project Overview")
make_table(
    ["Item", "Detail"],
    [
        ("Product Name",       "MIMS-CP Portal"),
        ("Full Form",          "Medical Information Management System + Customer Portal"),
        ("Domain",             "Pharmaceutical / Medical Information"),
        ("Purpose",            "End-to-end platform for handling medical inquiries from HCPs and patients, "
                               "with admin-controlled workflow, audit trail, and a future-facing customer portal (CP Portal)."),
        ("Compliance Target",  "21 CFR Part 11 — Electronic Records & Electronic Signatures"),
        ("Architecture",       "Monorepo — Two apps: MIMS (internal staff) + CP Portal (patient-facing, Phase 8)"),
        ("Tech Stack",         "React 19 + Node.js / Express + SQLite (PostgreSQL pre-deployment)"),
        ("Database",           "SQLite (development) — PostgreSQL migration planned pre-deployment"),
        ("Cloud / Staging",    "Not deployed (local development only — cloud staging in backlog)"),
        ("Repo Location",      "/Users/rohithkarne/MIMS-CP Portal/"),
    ],
    col_widths=[2.0, 4.8]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TEAM STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
section_heading("2", "Team Structure")
make_table(
    ["Role", "Person", "Ownership & Responsibilities"],
    [
        ("CPO — Product Owner",     "Rohith",       "Product strategy, all frontend development, feature decisions, sprint closures, client delivery, business & sales."),
        ("CTO — Technology Owner",  "Varun",        "Complete backend, infrastructure, security, cloud architecture, database design and all technology decisions."),
        ("Testing / QA Lead",       "Claude (AI)",  "Code review, QA testing, bug identification and reporting. Reports directly to Rohith (CPO)."),
    ],
    col_widths=[1.8, 1.2, 3.8]
)
para("Governance Rule: No sprint can be declared closed without explicit CPO (Rohith) approval.",
     bold=True, size=10, colour=RGBColor(0xC0, 0x39, 0x2B), space_before=2, space_after=6)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SPRINT ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
section_heading("3", "Sprint Roadmap")
make_table(
    ["Sprint", "Name", "Status", "Closed Date"],
    [
        ("Sprint 0",  "Project Scaffold",                         "✅ Complete",  "—"),
        ("Sprint 1",  "UI Overhaul",                              "✅ Complete",  "—"),
        ("Sprint 2",  "React Migration + Inbox UI",               "✅ CLOSED",    "11-Mar-2026"),
        ("Sprint 3",  "Admin Console",                            "✅ CLOSED",    "11-Mar-2026"),
        ("Sprint 4",  "New Case Form + Case Management Queues",   "📋 Pending",   "—"),
        ("Sprint 5",  "Case Query + Inbox Enhancements",          "📋 Backlog",   "—"),
        ("Sprint 6",  "Response Fulfillment + Correspondence",    "📋 Backlog",   "—"),
        ("Sprint 7",  "Transmissions + Browse Content + Analytics","📋 Backlog",  "—"),
        ("Sprint 8",  "CP Portal (Customer-Facing)",              "📋 Backlog",   "—"),
        ("Sprint 9",  "Integrations (Oracle Argus, Veeva)",       "📋 Backlog",   "—"),
    ],
    col_widths=[1.0, 2.8, 1.4, 1.6]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WHAT HAS BEEN BUILT
# ══════════════════════════════════════════════════════════════════════════════
section_heading("4", "Features Delivered (Sprint 0 – Sprint 3)")

sub_heading("4.1  Authentication System")
for item in [
    "User registration — name, email, password, role selection",
    "Login with JWT-based session management (8-hour token expiry)",
    "Password hashing via bcrypt (10 rounds — industry standard)",
    "Protected route guards on the frontend (unauthorised users redirected to login)",
    "Role-based access control (RBAC) middleware on the backend",
    "Login / logout audit trail per 21 CFR Part 11 requirements",
    "Failed login attempts captured with reason for compliance reporting",
]:
    bullet(item)

sub_heading("4.2  Dashboard")
for item in [
    "Personalised welcome banner with logged-in user's first name",
    "4 stat cards: Total Inquiries, Open Cases, Pending Review, Closed Today",
    "Quick Actions: Go to Inbox, Case Queue (stub), Manage Users (stub)",
    "Light/Dark theme toggle with localStorage persistence",
    "Sidebar collapse/expand state persisted across sessions",
]:
    bullet(item)

sub_heading("4.3  Inbox Module (Sprint 2 — 10 User Stories)")
make_table(
    ["Story", "Feature Delivered"],
    [
        ("US-01", "Inquiry list grouped by Today / Older"),
        ("US-02", "5 tabs: Inbox, Pending, Processed, Non-Processed, Outbox — with live counts"),
        ("US-03", "Real-time keyword search (filters by sender name and subject)"),
        ("US-04", "Lock / Unlock inquiries per individual user"),
        ("US-05", "Color-code inquiries — Red, Yellow, Green, Blue"),
        ("US-06", "Split-view detail panel (right side) — full email with From/To/Dates/Attachments"),
        ("US-07", "Sort toggle: Newest First (default) / Oldest First"),
        ("US-08", "Pagination — 50 records per page, Prev/Next navigation"),
        ("US-09", "Attachment count displayed in the detail view header"),
        ("US-10", "Timezone banner — 'Dates displayed in [TIMEZONE]'"),
    ],
    col_widths=[0.9, 5.9]
)
para("Data: 151 seeded inquiry records from backend. All inbox state (lock, color, tab) persists in localStorage per user.",
     size=9, colour=RGBColor(0x55, 0x55, 0x55), space_before=0, space_after=6)

sub_heading("4.4  Admin Console (Sprint 3 — Full Implementation)")
para("Accessible only to users with the admin role. Left-navigation with 30+ sections.", size=10, colour=DARK, space_before=0, space_after=4)
make_table(
    ["Section", "Features Delivered"],
    [
        ("Overview",               "Stats dashboard: organisations, workflows, sources, products, users, audit entries"),
        ("Sites Setup",            "Create organisations, add sites per org (name, country, primary flag), activate/deactivate"),
        ("Workflow Setup",         "Create and toggle workflow states active/inactive"),
        ("Source Types",           "Create and toggle source types: Email, Phone, Fax, CP Portal"),
        ("Product Dictionary",     "Create drug / trade names linked to organisations, activate/deactivate"),
        ("User Security Groups",   "4 roles × 9 modules permission matrix — toggle access per module per role"),
        ("User Configuration",     "Create users (admin-only), assign roles, view all system users"),
        ("Admin Audit Trail",      "All CREATE / UPDATE / DELETE / ESIG events — filterable by date, user, action — CSV export"),
        ("Login Audit Trail",      "Login / logout events per 21 CFR Part 11 — filterable by date, user — CSV export"),
        ("30+ Coming Soon Sections","Picklists, Email Accounts, Case Numbering, Data Protection, Transmission Rules, Product Groups, Field Setup, Case Form Definition, Custom Forms, Contact Repositories, Analytics, Integrations, SSP Setup, Help System, and more"),
    ],
    col_widths=[2.0, 4.8]
)

sub_heading("4.5  Electronic Signature (21 CFR Part 11)")
for item in [
    "Modal prompt appears before any critical admin action (org deactivation, permission changes, etc.)",
    "Requires current password (identity verification) + written reason (justification)",
    "Backend verifies password at moment of action — cannot be bypassed by frontend manipulation",
    "All verified actions written to audit trail with action type ESIG",
]:
    bullet(item)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — BACKEND API
# ══════════════════════════════════════════════════════════════════════════════
section_heading("5", "Backend API — Full Endpoint Inventory")

sub_heading("Auth  (/api/auth)")
make_table(
    ["Method", "Endpoint", "Purpose"],
    [
        ("POST", "/register", "Register new user — name, email, password, role"),
        ("POST", "/login",    "Authenticate user and issue JWT token"),
        ("GET",  "/me",       "Return current authenticated user details"),
        ("POST", "/logout",   "Record logout timestamp in login audit trail"),
    ],
    col_widths=[0.8, 1.4, 4.6]
)

sub_heading("Inbox  (/api/inbox)")
make_table(
    ["Method", "Endpoint", "Purpose"],
    [("GET", "/", "Return all 151 seeded inquiry records")],
    col_widths=[0.8, 1.4, 4.6]
)

sub_heading("Admin — Organisations & Sites  (/api/admin/orgs)")
make_table(
    ["Method", "Endpoint", "Purpose"],
    [
        ("GET",  "/",           "List all organisations"),
        ("POST", "/",           "Create new organisation"),
        ("PUT",  "/:id",        "Update organisation (name, is_active)"),
        ("GET",  "/:id/sites",  "List all sites for a given organisation"),
        ("POST", "/:id/sites",  "Create new site under an organisation"),
        ("PUT",  "/sites/:id",  "Update site (name, country, is_primary, is_active)"),
    ],
    col_widths=[0.8, 1.6, 4.4]
)

sub_heading("Admin — Configuration  (/api/admin/...)")
make_table(
    ["Endpoint", "Method(s)", "Purpose"],
    [
        ("/workflow-states",  "GET, POST, PUT", "Workflow state CRUD"),
        ("/source-types",     "GET, POST, PUT", "Source type CRUD"),
        ("/products",         "GET, POST, PUT", "Product / drug name CRUD"),
        ("/users",            "GET, POST",       "Admin user management — list and create"),
        ("/permissions",      "GET, PUT",        "Role-permission matrix — read and update"),
        ("/audit-logs",       "GET",             "Filtered admin audit trail (date, user, action)"),
        ("/login-audit",      "GET",             "Filtered login/logout history (date, user)"),
        ("/esig-verify",      "POST",            "Electronic signature verification — password + reason"),
    ],
    col_widths=[2.2, 1.5, 3.1]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
section_heading("6", "Database Schema (SQLite)")
make_table(
    ["Table", "Purpose", "Key Columns"],
    [
        ("users",            "All system users",                       "id, name, email, password (bcrypt), role, is_active, created_at"),
        ("sessions",         "JWT session tracking",                   "user_id, token, expires_at"),
        ("organisations",    "Customer organisations",                 "id, name, is_active, created_at"),
        ("sites",            "Sites under each organisation",          "id, org_id, name, country, is_primary, is_active"),
        ("workflow_states",  "Case workflow stages",                   "id, name, is_active"),
        ("source_types",     "Inquiry source channels",                "id, name, is_active"),
        ("products",         "Drug / trade name registry",             "id, trade_name, org_id, is_active"),
        ("audit_logs",       "All configuration change events",        "user_id, user_name, action, entity, entity_id, details, created_at"),
        ("login_audit",      "Login / logout event history",           "user_id, user_name, role, login_time, logout_time, status, fail_reason"),
        ("role_permissions", "Module access control per role",         "role, module, can_access"),
    ],
    col_widths=[1.5, 2.3, 3.0]
)

sub_heading("Roles & Modules")
para("Roles:  admin  |  agent  |  reviewer  |  content_manager", size=10, colour=DARK, space_before=0, space_after=4)
para("Modules:  inbox  |  case_mgmt  |  case_query  |  utilities  |  transmissions  |  browse_content  |  analytics  |  user_mgmt  |  admin_console", size=10, colour=DARK, space_before=0, space_after=4)
make_table(
    ["Role", "Modules With Access"],
    [
        ("admin",           "ALL 9 modules"),
        ("agent",           "inbox, case_mgmt, case_query, utilities, transmissions, browse_content"),
        ("reviewer",        "inbox, case_mgmt, case_query, browse_content, analytics"),
        ("content_manager", "browse_content only"),
    ],
    col_widths=[1.8, 5.0]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — 21 CFR PART 11 COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
section_heading("7", "21 CFR Part 11 Compliance Coverage")
make_table(
    ["Requirement", "Status", "Implementation"],
    [
        ("IMP-01 — Audit trail for all changes",          "✅ Done", "audit_logs table auto-populated on every admin action"),
        ("AUD-01 — Log all configuration changes",        "✅ Done", "All admin routes write to audit_logs on CREATE / UPDATE / DELETE"),
        ("AUD-02 — Login audit trail",                    "✅ Done", "login_audit table populated on every login attempt"),
        ("AUD-03 — Logout time recorded",                 "✅ Done", "POST /api/auth/logout updates logout_time in login_audit"),
        ("AUD-04 — Failed login attempts logged",         "✅ Done", "Failed logins recorded with reason field"),
        ("ESIG-01 — E-signature before critical actions", "✅ Done", "Modal appears before org/permission changes in Admin Console"),
        ("ESIG-02 — Signature verified at time of action","✅ Done", "Backend verifies password via /api/admin/esig-verify"),
        ("ESIG-03 — Signature written to audit trail",    "✅ Done", "ESIG entries written to audit_logs with action = ESIG"),
        ("ACC-01 — Role-based access control",            "✅ Done", "role_permissions table + requireRole() middleware on all protected routes"),
    ],
    col_widths=[2.6, 0.9, 3.3]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — FRONTEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
section_heading("8", "Frontend Architecture")

sub_heading("8.1  Page Routes")
make_table(
    ["Path", "Page / Component", "Auth Required", "Description"],
    [
        ("/",         "Redirect → /login",  "No",  "Default redirect to login"),
        ("/login",    "LoginPage",          "No",  "Sign In + Register tabs"),
        ("/dashboard","DashboardPage",      "Yes", "Main dashboard with stats and quick actions"),
        ("/inbox",    "InboxPage",          "Yes", "Full inbox module (10 user stories)"),
        ("/admin",    "AdminConsolePage",   "Yes (admin)", "Full admin console"),
        ("*",         "Redirect → /login", "—",   "Catch-all fallback"),
    ],
    col_widths=[1.2, 1.8, 1.2, 2.6]
)

sub_heading("8.2  Shared Components")
make_table(
    ["Component", "Purpose"],
    [
        ("Sidebar.jsx",       "Left navigation menu with section links and light/dark theme switcher"),
        ("Topbar.jsx",        "Top header bar with page title, user initials and logout button"),
        ("ProtectedRoute.jsx","Auth guard — checks JWT in context; redirects to login if missing"),
    ],
    col_widths=[1.8, 5.0]
)

sub_heading("8.3  Design System (index.css)")
make_table(
    ["Token", "Value", "Use"],
    [
        ("--primary",     "#1A5276", "Dark blue — primary actions, buttons, headings"),
        ("--accent",      "#17A589", "Teal — highlights, active states, accents"),
        ("--sidebar-bg",  "#12334F", "Dark sidebar background"),
        ("--bg",          "#F0F4F8", "Light page background"),
        ("--danger",      "#C0392B", "Errors, destructive actions"),
        ("--success",     "#1E8449", "Active status, confirmed states"),
        ("--warning",     "#D68910", "Warnings, locked states"),
    ],
    col_widths=[1.5, 1.3, 4.0]
)
para("Supports: Light mode and Dark mode via [data-theme='dark'] CSS variables.", size=10, colour=DARK, space_before=0, space_after=6)

sub_heading("8.4  Local Storage Keys")
make_table(
    ["Key", "Purpose", "Example Value"],
    [
        ("mims_user",                "Current logged-in user object",       '{ id, name, email, role }'),
        ("mims_token",               "JWT auth token",                      "eyJhbGciOiJIUzI1NiIs..."),
        ("mims_theme",               "Theme preference",                    "'light' or 'dark'"),
        ("mims_sidebar_collapsed",   "Sidebar collapse state",              "'true' or 'false'"),
        ("mims_inbox_{userId}",      "Per-user inbox state (lock + color)", "Array of inquiry objects"),
    ],
    col_widths=[2.0, 2.4, 2.4]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — SECURITY
# ══════════════════════════════════════════════════════════════════════════════
section_heading("9", "Security Implementation")
make_table(
    ["Control", "Implementation"],
    [
        ("Password Hashing",      "bcrypt with 10 salt rounds — no plaintext passwords stored at any point"),
        ("JWT Tokens",            "Signed with server secret — 8-hour expiry — sent as Bearer token in Authorization header"),
        ("Password Minimum",      "Minimum 8 characters enforced at registration and user creation"),
        ("RBAC Middleware",       "requireRole() middleware on all admin endpoints — role checked server-side"),
        ("E-Signature",           "Critical actions require re-authentication (password + reason) at time of action"),
        ("CORS",                  "Cross-origin policy configured via Express CORS middleware"),
        ("Soft Delete",           "Records set is_active = 0 rather than hard-deleted — full data integrity maintained"),
    ],
    col_widths=[1.8, 5.0]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — KEY FILE INVENTORY
# ══════════════════════════════════════════════════════════════════════════════
section_heading("10", "Key File Inventory")
make_table(
    ["File", "Location", "Purpose"],
    [
        ("App.jsx",              "frontend/src/",                     "React Router config — all route definitions"),
        ("main.jsx",             "frontend/src/",                     "React entry point"),
        ("index.css",            "frontend/src/",                     "Full design system — all CSS variables, components, layouts"),
        ("LoginPage.jsx",        "frontend/src/pages/",               "Sign In + Register page"),
        ("DashboardPage.jsx",    "frontend/src/pages/",               "Main dashboard"),
        ("InboxPage.jsx",        "frontend/src/pages/",               "Full inbox module — 10 user stories"),
        ("AdminConsolePage.jsx", "frontend/src/pages/",               "Admin Console — full Sprint 3 implementation"),
        ("Sidebar.jsx",          "frontend/src/components/",          "Navigation sidebar with theme switcher"),
        ("Topbar.jsx",           "frontend/src/components/",          "Header bar with user info and logout"),
        ("ProtectedRoute.jsx",   "frontend/src/components/",          "Auth guard wrapper"),
        ("AuthContext.jsx",      "frontend/src/context/",             "Global auth state — login, logout, token management"),
        ("server.js",            "backend/",                          "Express entry point — middleware and route mounts"),
        ("db.js",                "backend/database/",                 "SQLite schema, table creation, seed data"),
        ("auth.js (middleware)", "backend/middleware/",               "JWT verification + requireRole() RBAC"),
        ("authController.js",    "backend/controllers/",              "Login, register, logout logic"),
        ("auth.js (routes)",     "backend/routes/",                   "Auth API endpoints"),
        ("inbox.js",             "backend/routes/",                   "Inbox seed data endpoint — 151 records"),
        ("orgs.js",              "backend/routes/admin/",             "Organisation and site CRUD"),
        ("config.js",            "backend/routes/admin/",             "Workflows, sources, products, permissions, audit, users, ESIG"),
    ],
    col_widths=[2.0, 2.0, 2.8]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — DEFERRED BACKLOG
# ══════════════════════════════════════════════════════════════════════════════
section_heading("11", "Deferred Backlog (Not Yet Scheduled)")
make_table(
    ["Item", "Notes"],
    [
        ("Products: duplicate trade name check",       "Per-org uniqueness validation — deferred from Sprint 3"),
        ("User edit after creation",                   "Role/org changes post-creation — deferred from Sprint 3"),
        ("Audit: old value → new value on UPDATE",     "Stricter 21 CFR Part 11 compliance — deferred from Sprint 3"),
        ("PostgreSQL migration",                       "Replace SQLite before production deployment"),
        ("Email integration (Outlook / Gmail)",        "Planned for Phase 5–6 scope"),
        ("Cloud staging (Railway / Render)",           "No staging environment yet — all local development"),
        ("GDPR / HIPAA / DPPR compliance features",   "Phase TBD — requires dedicated scope discussion"),
        ("Oracle Argus integration",                   "Sprint 9"),
        ("Veeva CRM / Vault integration",              "Sprint 9"),
        ("CP Portal (customer-facing)",                "Sprint 8"),
        ("Admin force-unlock (permission-based)",      "Unscheduled"),
        ("3rd Case Management sub-tab name",           "'Team Cases' placeholder — Rohith to confirm final name"),
        ("Utilities module definition",                "Scope not yet defined"),
        ("Junior tester automation pipeline",          "Unscheduled"),
    ],
    col_widths=[2.8, 4.0]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — DEPENDENCIES
# ══════════════════════════════════════════════════════════════════════════════
section_heading("12", "Dependencies")

sub_heading("Frontend (npm)")
make_table(
    ["Package", "Version", "Purpose"],
    [
        ("react",                "^19.2.0",   "UI library"),
        ("react-dom",            "^19.2.0",   "DOM rendering"),
        ("react-router-dom",     "^7.13.1",   "Client-side routing"),
        ("vite",                 "^7.3.1",    "Build tool and dev server"),
        ("@vitejs/plugin-react", "^5.1.1",    "Vite React plugin"),
        ("eslint",               "^9.39.1",   "Linting"),
    ],
    col_widths=[2.2, 1.4, 3.2]
)

sub_heading("Backend (npm)")
make_table(
    ["Package", "Version", "Purpose"],
    [
        ("express",         "^4.18.2", "Web framework"),
        ("bcrypt",          "^5.1.1",  "Password hashing"),
        ("better-sqlite3",  "^9.4.3",  "SQLite database driver"),
        ("jsonwebtoken",    "^9.0.2",  "JWT creation and verification"),
        ("cors",            "^2.8.5",  "CORS middleware"),
        ("nodemon",         "^3.0.3",  "Auto-restart on file changes (dev only)"),
    ],
    col_widths=[2.2, 1.4, 3.2]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════
section_heading("13", "How to Run the System")
make_table(
    ["Step", "Command / Action"],
    [
        ("1. Start Backend",   "cd 'MIMS-CP Portal/mims/backend'  →  node server.js  →  Runs on http://localhost:3000"),
        ("2. Start Frontend",  "cd 'MIMS-CP Portal/mims/frontend'  →  npm run dev  →  Runs on http://localhost:5173"),
        ("3. Open Browser",    "Navigate to http://localhost:5173  (NOT port 3000)"),
        ("4. Register User",   "Use the Register tab on the login page to create the first account"),
        ("5. Admin Access",    "Register with role = admin to access the Admin Console at /admin"),
    ],
    col_widths=[1.6, 5.2]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — NEXT SPRINT (SPRINT 4)
# ══════════════════════════════════════════════════════════════════════════════
section_heading("14", "Next Sprint — Sprint 4")
make_table(
    ["Item", "Detail"],
    [
        ("Sprint Name",  "New Case Form + Case Management Queues"),
        ("Status",       "Scope definition in progress — awaiting CPO (Rohith) sign-off before kickoff"),
        ("Planned Scope","New Case creation form with all required fields; Case Management queues (My Cases, Team Cases, All Cases, Unassigned)"),
        ("Owner",        "Rohith (Frontend) + Varun (Backend)"),
        ("Testing",      "Claude (AI Testing Lead)"),
    ],
    col_widths=[1.8, 5.0]
)


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("— End of Handover Document —")
run.font.size      = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic         = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("MIMS-CP Portal  |  Sprint 3 Close  |  11 March 2026  |  Internal — Confidential")
run2.font.size      = Pt(9)
run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = "/Users/rohithkarne/MIMS-CP Portal/MIMS-CP Portal — PMO Handover Summary (Sprint 3 Close).docx"
doc.save(output_path)
print(f"Saved: {output_path}")
